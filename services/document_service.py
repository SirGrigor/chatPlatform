import json
import os
import shutil
import uuid

import chardet
from fastapi import UploadFile
from sqlalchemy import func
from db.models.document import Document
from db.session import DBSession

db_session = DBSession().get_db()


def extract_pdf_metadata(filepath: str) -> dict:
    import fitz  # PyMuPDF
    doc = fitz.open(filepath)
    metadata = doc.metadata
    doc.close()
    return metadata


def extract_docx_metadata(filepath: str) -> dict:
    from docx import Document
    doc = Document(filepath)
    core_properties = doc.core_properties

    metadata = {
        "author": core_properties.author,
        "title": core_properties.title,
    }
    return metadata


def read_pdf_content(filepath: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(filepath)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def read_word_content(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return text


def read_document_content(filepath: str) -> str:
    """
    Reads and returns the document content based on the file type.
    The function assumes that the file type has been validated and is one of the allowed types.
    """
    file_extension = filepath.split('.')[-1].lower()
    if file_extension in ['json', 'xml']:
        with open(filepath, "rb") as file:
            raw_data = file.read()
            encoding = chardet.detect(raw_data)['encoding']
            return raw_data.decode(encoding)
    elif file_extension == 'pdf':
        # Use PyMuPDF (fitz) to read PDF documents
        return read_pdf_content(filepath)
    elif file_extension == 'docx':
        # Use python-docx to read Word documents
        return read_word_content(filepath)
    else:
        raise ValueError("Unsupported file type for reading content.")


def save_document_file(file, file_path: str):
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)


class DocumentService:
    def __init__(self, db):
        self.db = db

    def create_document(self, course_id: int, filename: str, filepath: str, file_type: str,
                        document_content: str, document_metadata: dict) -> Document:
        document_metadata = json.dumps(document_metadata) if document_metadata else None
        db_document = Document(course_id=course_id, filename=filename,
                               filepath=filepath, file_type=file_type,
                               created_at=func.now(), updated_at=func.now(),
                               document_content=document_content, document_metadata=document_metadata)
        self.db.add(db_document)
        self.db.commit()
        self.db.refresh(db_document)
        return db_document

    def get_documents_for_course(self, course_id: int):
        return self.db.query(Document).filter(Document.course_id == course_id).all()

    def upload_document(self, course_id: int, file: UploadFile,
                        base_path: str = os.getenv("DOCUMENT_STORAGE_PATH", "/app/documents")):
        filename = file.filename
        file_extension = filename.split(".")[-1].lower()  # Ensure extension is in lowercase for comparison
        allowed_extensions = {"json", "xml", "pdf", "docx"}
        if file_extension not in allowed_extensions:
            raise ValueError("Unsupported file type.")

        file_path = os.path.join(base_path, f"{uuid.uuid4()}.{file_extension}")
        save_document_file(file, file_path)

        document_content = None
        document_metadata = None  # Initialize a variable to store metadata

        if file_extension in ["json", "xml"]:
            document_content = read_document_content(file_path)
        elif file_extension == "pdf":
            document_metadata = extract_pdf_metadata(file_path)
            document_content = read_document_content(file_path)
        elif file_extension == "docx":
            document_metadata = extract_docx_metadata(file_path)
            document_content = read_document_content(file_path)

        return self.create_document(course_id=course_id, filename=filename, filepath=file_path,
                                    file_type=file.content_type, document_content=document_content,
                                    document_metadata=document_metadata)

    def delete_document(self, document_id: int) -> bool:
        """
        Deletes a document by its ID.

        :param db: SQLAlchemy Session object.
        :param document_id: ID of the document to delete.
        :return: True if the document was deleted, False otherwise.
        """
        db_document = self.db.query(Document).filter(Document.id == document_id).first()
        if db_document:
            # Delete physical file
            if os.path.exists(db_document.filepath):
                os.remove(db_document.filepath)
            self.db.delete(db_document)
            self.db.commit()
            return True
        return False
